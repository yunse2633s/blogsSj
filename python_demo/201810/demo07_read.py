#!coding:utf-8
#!/usr/lib/python3

#读word文件

import docx 


docStr=docx.Document('G:/work/python/201810/demo07.docx')

print(type(docStr)) # <class 'docx.document.Document'>
for x in docStr.paragraphs:
	#类型判断，若有子结构继续
	print(type(x), x.text)

for y in docStr.tables:
	#类型判断，若有子结构继续
	# print(dir(y))
	for z in y.rows:
		# print(dir(z))
		for za in z.cells:
			# print(za.paragraphs, za.tables)
			print(za.text)
			# for zb in za.paragraphs:
			